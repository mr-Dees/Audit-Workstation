"""Тест-страж паритета подписей/полноты нарушения между DOCX/MD/TXT (#32).

Блочная модель: reference-нарушение со ВСЕМИ 10 полями enabled+filled
(текст-блоки с уникальными маркерами; в additionalContent — картинка-черновик
с пустым url и текст-блок; в codeMining — блок-таблица) прогоняется через все
три форматтера:

- **label-parity** — метка каждого поля с ``labeled=True``
  (``violation_fields.LABELS``) обязана присутствовать в КАЖДОМ выводе, а
  метки полей с ``labeled=False`` — отсутствовать в КАЖДОМ;
- **value-parity** — уникальный маркер каждого поля доходит до каждого
  вывода (проверяем доходимость ЗНАЧЕНИЯ, не только метки); плейсхолдер
  картинки «Изображение: {filename}» одинаков во всех трёх форматтерах;
- **order-parity** — пользовательский fieldOrder уважают все три формата.

Источник ожидаемых меток — ``app.domains.acts.violation_fields.LABELS``.
Если кто-то откатит метку форматтера на старую — эти тесты упадут.
"""
from __future__ import annotations

import pytest
from docx import Document

from app.domains.acts.formatters.docx.builders.violation import build_violation
from app.domains.acts.formatters.markdown_formatter import MarkdownFormatter
from app.domains.acts.formatters.text_formatter import TextFormatter
from app.domains.acts.schemas.act_content import ViolationSchema
from app.domains.acts.settings import ActsSettings
from app.domains.acts.violation_fields import (
    FIELD_BY_KEY,
    LABELS,
    VIOLATION_FIELD_KEYS,
)


def _md() -> MarkdownFormatter:
    return MarkdownFormatter(settings=None, acts_settings=ActsSettings())


def _txt() -> TextFormatter:
    return TextFormatter(settings=None, acts_settings=ActsSettings())


def _docx_text(violation: ViolationSchema) -> str:
    """Рендерит нарушение в свежий Document и возвращает весь текст (абзацы+таблицы)."""
    doc = Document()
    build_violation(doc, violation)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# Маркер значения на каждое поле реестра (текст-блоком).
FIELD_MARKERS = {key: f"МАРКЕР_{key.upper()}" for key in VIOLATION_FIELD_KEYS}
IMAGE_FILENAME = "МАРКЕР_FILE.png"
TABLE_CELL_MARKER = "МАРКЕР_ЯЧЕЙКА"


def _text_block(content: str, bid: str) -> dict:
    return {"id": bid, "type": "text", "content": content}


def _reference_violation(field_order=None) -> ViolationSchema:
    payload = {"id": "v1", "nodeId": "n1", "fieldOrder": field_order}
    for i, key in enumerate(VIOLATION_FIELD_KEYS):
        payload[key] = {
            "enabled": True,
            "blocks": [_text_block(FIELD_MARKERS[key], f"text_{i}")],
        }
    # additionalContent: + картинка-черновик (пустой url → одинаковый
    # текст-плейсхолдер во всех трёх форматтерах).
    payload["additionalContent"]["blocks"].append({
        "id": "image_x", "type": "image",
        "url": "", "caption": "", "filename": IMAGE_FILENAME, "width": 0,
    })
    # codeMining: + блок-таблица.
    payload["codeMining"]["blocks"].append({
        "id": "table_x", "type": "table",
        "table": {"grid": [[{"content": TABLE_CELL_MARKER}]], "colWidths": [100]},
    })
    return ViolationSchema.model_validate(payload)


def _all_outputs(violation: ViolationSchema) -> dict[str, str]:
    data = violation.model_dump()
    return {
        "docx": _docx_text(violation),
        "md": _md()._format_violation(data),
        "txt": _txt()._format_violation(data),
    }


@pytest.mark.parametrize("fmt", ["docx", "md", "txt"])
def test_labels_parity(fmt):
    """Метка КАЖДОГО поля с labeled=True присутствует в выводе форматтера."""
    out = _all_outputs(_reference_violation())[fmt]
    for key in VIOLATION_FIELD_KEYS:
        if not FIELD_BY_KEY[key].labeled:
            continue
        assert LABELS[key] in out, f"{fmt}: метка {LABELS[key]!r} не дошла до вывода"


@pytest.mark.parametrize("fmt", ["docx", "md", "txt"])
def test_unlabeled_fields_have_no_label_in_any_format(fmt):
    """Метки полей с labeled=False не выводит НИ ОДИН формат (решение владельца).

    Симметрично label-parity: если кто-то вернёт метку одному форматтеру —
    паритет «без заголовка» сломается молча, этот тест поймает.
    """
    out = _all_outputs(_reference_violation())[fmt]
    for key in VIOLATION_FIELD_KEYS:
        if FIELD_BY_KEY[key].labeled:
            continue
        assert LABELS[key] not in out, (
            f"{fmt}: поле {key} выводится без метки — {LABELS[key]!r} лишняя"
        )


@pytest.mark.parametrize("fmt", ["docx", "md", "txt"])
def test_values_parity(fmt):
    """Маркер значения каждого поля доходит до вывода каждого форматтера."""
    out = _all_outputs(_reference_violation())[fmt]
    for key, marker in FIELD_MARKERS.items():
        assert marker in out, f"{fmt}: значение поля {key} не дошло до вывода"
    assert TABLE_CELL_MARKER in out, f"{fmt}: ячейка блока-таблицы не дошла"
    # Плейсхолдер картинки-черновика формат-специфичен (#16): MD — курсивное
    # имя файла, DOCX/TXT — строка «Изображение: {filename}».
    if fmt == "md":
        assert f"*{IMAGE_FILENAME}*" in out, f"{fmt}: плейсхолдер картинки не дошёл"
    else:
        assert f"Изображение: {IMAGE_FILENAME}" in out, f"{fmt}: плейсхолдер картинки не дошёл"


@pytest.mark.parametrize("fmt", ["docx", "md", "txt"])
def test_field_order_parity(fmt):
    """Пользовательский fieldOrder уважают все три формата."""
    order = list(VIOLATION_FIELD_KEYS)
    order.remove("responsible")
    order.insert(0, "responsible")
    out = _all_outputs(_reference_violation(field_order=order))[fmt]
    assert out.index(FIELD_MARKERS["responsible"]) < out.index(FIELD_MARKERS["violated"]), (
        f"{fmt}: поле, поднятое fieldOrder наверх, обязано рендериться первым"
    )
